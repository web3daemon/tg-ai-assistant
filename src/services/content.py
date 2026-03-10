import base64
import io
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from openpyxl import load_workbook

from src.config import settings


class ContentExtractionError(Exception):
    pass


@dataclass(slots=True)
class ExtractedDocument:
    file_name: str
    media_type: str
    extracted_text: str
    user_text: str

    @property
    def log_content(self) -> str:
        base = f"[document:{self.media_type}] {self.file_name}"
        if self.user_text:
            base += f"\nКомментарий: {self.user_text}"
        return base

    @property
    def model_text(self) -> str:
        prompt = self.user_text or "Проанализируй документ."
        return (
            f"{prompt}\n\n"
            f"Имя файла: {self.file_name}\n"
            f"Тип файла: {self.media_type}\n\n"
            "Содержимое документа:\n"
            f"{self.extracted_text}"
        )


@dataclass(slots=True)
class ExtractedImage:
    file_name: str
    mime_type: str
    base64_data: str
    user_text: str

    @property
    def log_content(self) -> str:
        base = f"[image:{self.mime_type}] {self.file_name}"
        if self.user_text:
            base += f"\nКомментарий: {self.user_text}"
        return base

    @property
    def prompt_text(self) -> str:
        return self.user_text or "Опиши и проанализируй изображение."


def ensure_supported_file_size(file_size: int | None) -> None:
    if file_size is None:
        return

    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if file_size > max_bytes:
        raise ContentExtractionError(
            f"Файл слишком большой. Лимит: {settings.max_file_size_mb} МБ."
        )


def extract_document(file_name: str, data: bytes, user_text: str) -> ExtractedDocument:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".txt":
        text = _extract_txt(data)
        media_type = "txt"
    elif suffix == ".pdf":
        text = _extract_pdf(data)
        media_type = "pdf"
    elif suffix == ".docx":
        text = _extract_docx(data)
        media_type = "docx"
    elif suffix == ".xlsx":
        text = _extract_xlsx(data)
        media_type = "xlsx"
    else:
        raise ContentExtractionError(
            "Неподдерживаемый документ. Поддерживаются: .txt, .pdf, .docx, .xlsx."
        )

    text = _trim_text(text)
    if not text:
        raise ContentExtractionError("Не удалось извлечь текст из документа.")

    return ExtractedDocument(
        file_name=file_name,
        media_type=media_type,
        extracted_text=text,
        user_text=user_text,
    )


def extract_image(file_name: str, mime_type: str | None, data: bytes, user_text: str) -> ExtractedImage:
    suffix = Path(file_name).suffix.lower()
    resolved_mime_type = mime_type or _guess_image_mime_type(suffix)
    if resolved_mime_type not in {"image/jpeg", "image/png", "image/webp"}:
        raise ContentExtractionError("Неподдерживаемый формат изображения.")

    return ExtractedImage(
        file_name=file_name,
        mime_type=resolved_mime_type,
        base64_data=base64.b64encode(data).decode("ascii"),
        user_text=user_text,
    )


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ContentExtractionError("Не удалось декодировать текстовый файл.")


def _extract_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as document:
        for page in document:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    document = DocxDocument(io.BytesIO(data))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(lines)


def _extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Лист: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _trim_text(text: str) -> str:
    normalized = text.strip()
    if len(normalized) <= settings.max_extracted_text_chars:
        return normalized
    return normalized[: settings.max_extracted_text_chars].rstrip() + "\n\n[Обрезано по лимиту]"


def _guess_image_mime_type(suffix: str) -> str:
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".png":
        return "image/png"
    if suffix == ".webp":
        return "image/webp"
    return "application/octet-stream"
